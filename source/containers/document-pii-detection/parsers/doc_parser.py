
import docx
from .parser import BaseParser

class DocParser(BaseParser):
    def __init__(self, s3_client):
        super().__init__(s3_client=s3_client)

    def parse_file(self, doc_path):
        """
        Extracts text from a doc file and returns a string of content.
        """

        doc = docx.Document(doc_path)
        file_content = ""
        for para in doc.paragraphs:
            file_content += para.text + "\n"
        
        prev_paragraph = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text != prev_paragraph:
                            file_content += paragraph.text + "\n"
                            prev_paragraph = paragraph.text

        return [file_content]
